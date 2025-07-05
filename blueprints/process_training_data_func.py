import io
import json
import logging
import unicodedata
import azure.functions as func
from pathlib import Path

from configs.openai_client import AzureOpenAIClient
from models.DocumentMetadata import DocumentMetadata
from constants import BLOB_CONTAINER_NAME

from langchain_experimental.text_splitter import SemanticChunker
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

from configs.settings import embeddings, vector_store
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

process_data_bp = func.Blueprint()

@process_data_bp.function_name(name="process-training-data")
@process_data_bp.blob_trigger(
    arg_name="blob",
    path=f"{BLOB_CONTAINER_NAME}/{{name}}",
    connection="AzureWebJobsStorage"
)
def process_raw_training_data(blob: func.InputStream):
    logging.info(f"Processando arquivo: {blob.name}")

    try:
        if not blob.name:
            logging.error("Blob name is None.")
            return

        blob_bytes = blob.read()
        ext = Path(blob.name).suffix.lower()

        if ext not in ['.pdf', '.docx', '.txt', '.md']:
            logging.error(f"Tipo de arquivo {ext} não suportado.")
            return

        full_content = get_blob_content(ext=ext, blob_bytes=blob_bytes)
        chunk_size = 512 # https://arxiv.org/pdf/2407.01219

        semantic_chunker = SemanticChunker(
            embeddings,
            breakpoint_threshold_type="percentile",
            breakpoint_threshold_amount=95.0,
            min_chunk_size=chunk_size
        )
        semantic_docs = semantic_chunker.create_documents([full_content])

        class_code = Path(blob.name).parent.name
        file_name = Path(blob.name).name

        enriched_docs = []
        for doc in semantic_docs:
            data = extract_metadata(doc.page_content)
            doc.metadata = {
                "file_id": file_name,
                "category": data.category,
                "subcategory": data.subcategory,
                "class_code": class_code
            }
            enriched_docs.append(doc)

        char_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, 
            chunk_overlap=60, # ~12% https://learn.microsoft.com/en-us/azure/search/vector-search-how-to-chunk-documents#content-overlap-considerations
            length_function=token_count,
            separators=[
                "\n\n",
                "\n",
                " ",
                ".",
                ",",
                "\u200b",  # Zero-width space
                "\uff0c",  # Fullwidth comma
                "\u3001",  # Ideographic comma
                "\uff0e",  # Fullwidth full stop
                "\u3002",  # Ideographic full stop
                "",
            ],
            is_separator_regex=False
        )
        final_docs = []
        for doc in enriched_docs:
            sub_chunks = char_splitter.split_text(doc.page_content)
            logging.debug(f"Bloco semântico com metadata {doc.metadata} gerou {len(sub_chunks)} sub-chunks.")
            for chunk in sub_chunks:
                cleared_text = chunk.replace("\n", " ").replace("\r", " ").strip()
                final_docs.append(
                    LangchainDocument(page_content=cleared_text, metadata=doc.metadata)
                )

        # 5) Armazena os chunks finais no vector store
        vector_store.add_documents(final_docs)
        logging.info(f"{len(final_docs)} chunks inseridos no vector store.")

    except Exception as e:
        logging.error(f"Erro ao processar o arquivo {blob.name}: {str(e)}")


def extract_metadata(text: str) -> DocumentMetadata:
    try:
        prompt = (
            "Analise o seguinte texto e identifique os seguintes metadados:\n"
            "- Tags relevantes\n"
            "- Categoria principal - o assunto principal do texto\n"
            "- Subcategoria - outros temas que representa o texto\n"
            f"Texto: {text}\n"
            "Responda no formato JSON: {\"tags\": [\"...\"], \"category\": \"...\", \"subcategory\": \"...\"}"
        )
        response_text = AzureOpenAIClient.create_completion_json(
            prompt=prompt, max_tokens=300, temperature=0.6
        )
        if not response_text:
            return DocumentMetadata(text=text, category="Outros", subcategory="Outros", tags=[])

        data = json.loads(response_text)
        metadata = DocumentMetadata(**data)
        metadata.text = text
        return metadata

    except Exception as e:
        logging.error(f"Erro ao extrair metadados: {str(e)}")
        return DocumentMetadata(text=text, category="Outros", subcategory="Outros", tags=[])


def get_blob_content(ext: str, blob_bytes: bytes) -> str:
    if ext in ['.txt', '.md']:
        decoded = blob_bytes.decode('utf-8', errors='ignore')
        return clean_utf8_text(decoded)

    elif ext == '.pdf':
        reader = PdfReader(io.BytesIO(blob_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return clean_utf8_text(text)

    elif ext == '.docx':
        doc = DocxDocument(io.BytesIO(blob_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)
        return clean_utf8_text(text)

    else:
        logging.error(f"Tipo de arquivo {ext} não suportado.")
        raise ValueError(f"Unsupported file type: {ext}")

def clean_utf8_text(raw_text: str) -> str:
    # Normaliza e remove caracteres de controle invisíveis
    text = unicodedata.normalize("NFKC", raw_text)
    text = ''.join(c for c in text if unicodedata.category(c)[0] != 'C')
    return text.strip()

def token_count(input_string) -> int:
    import tiktoken

    encoding = tiktoken.get_encoding("cl100k_base")
    tokens = encoding.encode(input_string)
    token_count = len(tokens)
    return token_count